#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
from docx import Document

def files_folder(path):
    if type(path) != str:
        return "Erro, argument path != string"
    files = os.listdir(path)
    documents = {}
    for file in files:
        text = Document(path+file)
        full_text = ''
        for p in text.paragraphs: 
            full_text = full_text+p.text
        documents[file] = full_text            
    return documents

def file(name):    
    if type(name) != str:
        return "Erro, argument name != string"
    text = Document(name)
    full_text = ''
    for p in text.paragraphs: 
        full_text = full_text+p.text
    return full_text